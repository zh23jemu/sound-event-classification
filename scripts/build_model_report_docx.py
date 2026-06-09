from __future__ import annotations

import re
import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]


def add_formatted_text(paragraph, text: str) -> None:
    """向段落中写入带基础 Markdown 行内格式的文本。

    当前报告主要使用反引号标记文件名、指标名和类别名。Word 版中将这些内容设置
    为等宽字体，便于读者区分代码路径、配置文件和类别标签。
    """

    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part[1:-1] if part.startswith("`") and part.endswith("`") else part)
        if part.startswith("`") and part.endswith("`"):
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    """设置表格单元格文本，并统一表格字号。"""

    cell.text = ""
    paragraph = cell.paragraphs[0]
    add_formatted_text(paragraph, text)
    for run in paragraph.runs:
        run.bold = bold
        run.font.size = Pt(9)


def add_markdown_table(document: Document, lines: list[str]) -> None:
    """把 Markdown 表格转换成 Word 表格。

    这里只处理本项目报告中使用的标准管道表。表格采用固定内容结构，避免把普通
    段落误识别为表格。
    """

    rows = []
    for line in lines:
        stripped = line.strip().strip("|")
        cells = [cell.strip() for cell in stripped.split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        rows.append(cells)

    if not rows:
        return

    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = True

    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            set_cell_text(table.cell(row_index, col_index), value, bold=row_index == 0)

    document.add_paragraph()


def add_image(document: Document, alt_text: str, image_path: Path) -> None:
    """插入报告图片，并添加简短图题。"""

    if not image_path.exists():
        paragraph = document.add_paragraph(style="Intense Quote")
        paragraph.add_run(f"图片缺失：{image_path}").italic = True
        return

    document.add_picture(str(image_path), width=Inches(6.2))
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(alt_text)
    run.italic = True
    run.font.size = Pt(9)


def apply_document_styles(document: Document) -> None:
    """设置 Word 报告的基础样式。"""

    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color in [
        ("Heading 1", 18, RGBColor(31, 78, 121)),
        ("Heading 2", 14, RGBColor(31, 78, 121)),
        ("Heading 3", 12, RGBColor(47, 84, 150)),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)


def build_docx(markdown_path: Path, output_path: Path) -> None:
    """读取模型报告 Markdown，并生成 Word 初稿。"""

    document = Document()
    apply_document_styles(document)

    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        if not line:
            index += 1
            continue

        if line.startswith("# "):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(line[2:].strip())
            run.bold = True
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(31, 78, 121)
            index += 1
            continue

        if line.startswith("## "):
            document.add_heading(line[3:].strip(), level=1)
            index += 1
            continue

        if line.startswith("### "):
            document.add_heading(line[4:].strip(), level=2)
            index += 1
            continue

        if line.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_markdown_table(document, table_lines)
            continue

        image_match = re.fullmatch(r"!\[(.+)]\((.+)\)", line)
        if image_match:
            alt_text = image_match.group(1)
            image_path = PROJECT_ROOT / image_match.group(2)
            add_image(document, alt_text, image_path)
            index += 1
            continue

        if line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_formatted_text(paragraph, line[2:].strip())
            index += 1
            continue

        paragraph = document.add_paragraph()
        add_formatted_text(paragraph, line)
        index += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="把项目 Markdown 报告转换为 Word 文档")
    parser.add_argument("--input", default="模型报告.md", help="输入 Markdown 文件")
    parser.add_argument("--output", default="模型报告.docx", help="输出 Word 文件")
    args = parser.parse_args()

    output_path = PROJECT_ROOT / args.output
    build_docx(PROJECT_ROOT / args.input, output_path)
    print(f"已生成：{output_path.name}")


if __name__ == "__main__":
    main()
